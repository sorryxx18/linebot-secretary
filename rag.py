from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import sqlite3
import subprocess
import tempfile
from dataclasses import dataclass
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Iterable

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = Path(os.getenv("DATA_DIR", BASE_DIR / "data"))
RAW_DIR = DATA_DIR / "raw"
EXTRACTED_DIR = DATA_DIR / "extracted"
DB_PATH = DATA_DIR / "index.sqlite3"
DRIVE_MANIFEST_PATH = DATA_DIR / "drive_manifest.json"

SUPPORTED_SUFFIXES = {".txt", ".md", ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".xlsm", ".ods", ".csv"}

SENSITIVE_TERMS = [
    "名冊", "女性同仁", "證照人員", "毒性化學", "毒化物", "公共危險物品", "高壓氣體",
    "水源", "消防通道", "狹小巷道", "車輛配置", "上班日", "直升機起降場", "救護車清冊",
]

KNOWN_TERMS = [
    "水源", "缺乏", "匱乏", "替代水源", "蓄水池", "池塘", "河川", "消防栓",
    "消防通道", "狹小巷道", "長照", "老福", "古蹟", "歷史建築", "公共危險物品",
    "高壓氣體", "毒性化學", "毒化物", "義消", "救護車", "EMT", "無人機", "直升機",
    "大高空", "救災車輛", "山坡地", "老舊聚落", "防災公園", "火災統計",
    "火警", "火災", "搶救", "出勤", "消防安全", "列管", "場所", "罰鍰",
    "古蹟", "歷史建築", "住警器", "容留", "防災", "證照", "EMT", "OHCA",
    "T-CERT", "搜救", "山域", "水域", "夜市", "宿舍", "駐地", "器材",
    "信義區", "大安區", "南港區", "莊敬", "信義分隊", "安和", "舊莊",
]

SYNONYMS = {
    "匱乏": ["缺乏", "不足"],
    "缺乏": ["匱乏", "不足"],
    "水源匱乏": ["水源缺乏"],
    "水源不足": ["水源缺乏"],
    "火警": ["火災"],
    "火災": ["火警"],
    "救護": ["救護車"],
    "義消": ["義勇消防"],
}

@dataclass
class SearchHit:
    file_name: str
    title: str
    content: str
    score: float
    page: str | None = None
    sheet: str | None = None


def ensure_dirs() -> None:
    RAW_DIR.mkdir(parents=True, exist_ok=True)
    EXTRACTED_DIR.mkdir(parents=True, exist_ok=True)


def taiwan_minguo_date() -> str:
    now = datetime.now(timezone(timedelta(hours=8)))
    return f"民國 {now.year - 1911} 年 {now.month} 月 {now.day} 日"


def init_db() -> None:
    ensure_dirs()
    with sqlite3.connect(DB_PATH) as con:
        con.execute(
            """
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY,
                path TEXT UNIQUE NOT NULL,
                file_name TEXT NOT NULL,
                sha256 TEXT NOT NULL,
                mtime REAL NOT NULL,
                extracted_at TEXT NOT NULL
            )
            """
        )
        con.execute(
            """
            CREATE TABLE IF NOT EXISTS chunks (
                id INTEGER PRIMARY KEY,
                document_id INTEGER NOT NULL,
                chunk_index INTEGER NOT NULL,
                title TEXT,
                content TEXT NOT NULL,
                page TEXT,
                sheet TEXT,
                FOREIGN KEY(document_id) REFERENCES documents(id)
            )
            """
        )
        con.execute(
            """
            CREATE VIRTUAL TABLE IF NOT EXISTS chunks_fts USING fts5(
                file_name, title, content,
                tokenize='unicode61 remove_diacritics 2'
            )
            """
        )
        con.execute("CREATE INDEX IF NOT EXISTS idx_chunks_doc ON chunks(document_id)")
        con.commit()


def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(path: Path) -> str:
    import fitz
    parts = []
    doc = fitz.open(path)
    for idx, page in enumerate(doc, start=1):
        txt = page.get_text("text") or ""
        if txt.strip():
            parts.append(f"[第 {idx} 頁]\n{txt}")
    return clean_text("\n\n".join(parts))


def extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for ti, table in enumerate(doc.tables, start=1):
        parts.append(f"[表格 {ti}]")
        for row in table.rows:
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return clean_text("\n".join(parts))


def extract_doc_with_textutil(path: Path) -> str:
    """Extract .doc text — mammoth (cross-platform) with macOS textutil as fallback."""
    try:
        import mammoth
        with path.open("rb") as f:
            result = mammoth.extract_raw_text(f)
        if result.value.strip():
            return clean_text(result.value)
    except Exception:
        pass
    # macOS fallback
    import platform
    if platform.system() == "Darwin":
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / (path.stem + ".txt")
            subprocess.run(["textutil", "-convert", "txt", "-output", str(out), str(path)], check=False, timeout=60)
            if out.exists():
                return clean_text(out.read_text(errors="ignore"))
    return ""


def extract_spreadsheet(path: Path) -> str:
    import pandas as pd
    ext = path.suffix.lower()
    engine = None
    if ext in {".xlsx", ".xlsm"}:
        engine = "openpyxl"
    elif ext == ".xls":
        engine = "xlrd"
    elif ext == ".ods":
        engine = "odf"
    sheets = pd.read_excel(path, sheet_name=None, dtype=str, engine=engine, header=None)
    parts = []
    for sheet_name, df in sheets.items():
        df = df.fillna("")
        rows_text = []
        for _, row in df.iterrows():
            vals = [str(v).strip() for v in row.tolist() if str(v).strip()]
            if vals:
                rows_text.append("  ".join(vals))
        if rows_text:
            parts.append(f"[工作表：{sheet_name}]")
            parts.extend(rows_text)
    return clean_text("\n".join(parts))


def extract_csv(path: Path) -> str:
    rows = []
    with path.open("r", encoding="utf-8-sig", errors="ignore", newline="") as f:
        for row in csv.reader(f):
            if any(cell.strip() for cell in row):
                rows.append("\t".join(row))
    return clean_text("\n".join(rows))


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".txt", ".md"}:
        return clean_text(path.read_text(encoding="utf-8", errors="ignore"))
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".doc":
        return extract_doc_with_textutil(path)
    if ext in {".xlsx", ".xls", ".xlsm", ".ods"}:
        return extract_spreadsheet(path)
    if ext == ".csv":
        return extract_csv(path)
    return ""


def chunk_text(text: str, max_chars: int = 1200, overlap: int = 0) -> list[str]:
    text = clean_text(text)
    if not text:
        return []
    paras = [p.strip() for p in re.split(r"\n\s*\n|(?<=。)\s*", text) if p.strip()]
    chunks, cur = [], ""
    for p in paras:
        if len(cur) + len(p) + 2 <= max_chars:
            cur = (cur + "\n" + p).strip()
        else:
            if cur:
                chunks.append(cur)
            if len(p) > max_chars:
                for i in range(0, len(p), max_chars - overlap):
                    chunks.append(p[i:i+max_chars])
                cur = ""
            else:
                cur = p
    if cur:
        chunks.append(cur)
    # add small overlap from previous chunk for continuity
    out = []
    prev_tail = ""
    for c in chunks:
        out.append((prev_tail + "\n" + c).strip() if prev_tail else c)
        prev_tail = c[-overlap:]
    return out


def iter_files() -> Iterable[Path]:
    ensure_dirs()
    for p in RAW_DIR.rglob("*"):
        if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES and not p.name.startswith("~$"):
            yield p


def index_file(path: Path, force: bool = False) -> tuple[bool, str]:
    init_db()
    sha = file_sha256(path)
    mtime = path.stat().st_mtime
    path_s = str(path)
    with sqlite3.connect(DB_PATH) as con:
        row = con.execute("SELECT id, sha256, mtime FROM documents WHERE path=?", (path_s,)).fetchone()
        if row and row[1] == sha and not force:
            return False, f"skip unchanged: {path.name}"
        text = extract_text(path)
        if not text:
            return False, f"no text extracted: {path.name}"
        extracted_path = EXTRACTED_DIR / f"{path.name}.txt"
        extracted_path.write_text(text, encoding="utf-8")
        if row:
            doc_id = row[0]
            con.execute("DELETE FROM chunks_fts WHERE rowid IN (SELECT id FROM chunks WHERE document_id=?)", (doc_id,))
            con.execute("DELETE FROM chunks WHERE document_id=?", (doc_id,))
            con.execute("UPDATE documents SET sha256=?, mtime=?, extracted_at=? WHERE id=?", (sha, mtime, datetime.now().isoformat(), doc_id))
        else:
            cur = con.execute(
                "INSERT INTO documents(path, file_name, sha256, mtime, extracted_at) VALUES(?,?,?,?,?)",
                (path_s, path.name, sha, mtime, datetime.now().isoformat()),
            )
            doc_id = cur.lastrowid
        for i, chunk in enumerate(chunk_text(text)):
            title = path.stem
            cur = con.execute(
                "INSERT INTO chunks(document_id, chunk_index, title, content) VALUES(?,?,?,?)",
                (doc_id, i, title, chunk),
            )
            chunk_id = cur.lastrowid
            con.execute(
                "INSERT INTO chunks_fts(rowid, file_name, title, content) VALUES(?,?,?,?)",
                (chunk_id, path.name, title, chunk),
            )
        con.commit()
    return True, f"indexed: {path.name}"


def reindex_all(force: bool = False) -> dict:
    init_db()
    results = []
    for p in iter_files():
        try:
            results.append(index_file(p, force=force)[1])
        except Exception as e:
            results.append(f"ERROR {p.name}: {e}")
    return stats() | {"results": results}


def stats() -> dict:
    init_db()
    with sqlite3.connect(DB_PATH) as con:
        docs = con.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
        chunks = con.execute("SELECT COUNT(*) FROM chunks").fetchone()[0]
    raw_files = sum(1 for _ in iter_files())
    return {"raw_files": raw_files, "documents": docs, "chunks": chunks, "db_path": str(DB_PATH)}


def expanded_terms(query: str) -> list[str]:
    terms: list[str] = []
    for t in KNOWN_TERMS:
        if t in query and t not in terms:
            terms.append(t)
    for k, vals in SYNONYMS.items():
        if k in query:
            for v in vals:
                if v not in terms:
                    terms.append(v)
    # Keep Latin/alphanumeric terms, but skip standalone Taiwan year numbers (115, 114...)
    # which match too many documents and dilute search quality
    for t in re.findall(r"[A-Za-z0-9]{2,}", query):
        if t.isdigit() and 100 <= int(t) <= 200:
            continue
        if t not in terms:
            terms.append(t)
    chinese = re.sub(r"[^\u4e00-\u9fff]", "", query)
    if not terms:
        for n in (4, 3, 2):
            for i in range(0, max(0, len(chinese) - n + 1)):
                gram = chinese[i:i+n]
                if gram in {"什麼", "哪些", "有什", "麼水", "地區", "資料", "嗎"}:
                    continue
                if any(gram in kt or kt in gram for kt in KNOWN_TERMS) and gram not in terms:
                    terms.append(gram)
    if not terms and chinese:
        terms.append(chinese)
    return terms[:20]


def normalize_query(query: str) -> str:
    terms = expanded_terms(query)
    return " OR ".join(terms[:12]) if terms else query


def search(query: str, limit: int = 8) -> list[SearchHit]:
    init_db()
    hits: list[SearchHit] = []
    seen_docs: set[str] = set()  # track by file_name to dedupe per-document
    seen_ids: set[int] = set()
    with sqlite3.connect(DB_PATH) as con:
        con.row_factory = sqlite3.Row

        def _add_hit(r, score: float) -> bool:
            cid = r["id"] if "id" in r.keys() else r[0]
            fname = r["file_name"]
            if cid in seen_ids:
                return False
            seen_ids.add(cid)
            # Allow at most 2 chunks per document to avoid big files dominating
            doc_count = sum(1 for h in hits if h.file_name == fname)
            if doc_count >= 2:
                return False
            page = r["page"] if "page" in r.keys() else None
            sheet = r["sheet"] if "sheet" in r.keys() else None
            hits.append(SearchHit(fname, r["title"] or "", r["content"], score, page, sheet))
            return True

        fts_q = normalize_query(query)
        try:
            rows = con.execute(
                """
                SELECT c.id, d.file_name, c.title, c.content, c.page, c.sheet, bm25(chunks_fts) AS score
                FROM chunks_fts
                JOIN chunks c ON chunks_fts.rowid = c.id
                JOIN documents d ON c.document_id = d.id
                WHERE chunks_fts MATCH ?
                ORDER BY score LIMIT ?
                """,
                (fts_q, limit * 4),  # fetch more, then dedupe per-doc
            ).fetchall()
        except sqlite3.OperationalError:
            rows = []
        for r in rows:
            if len(hits) >= limit:
                break
            _add_hit(r, float(r["score"]))

        # Step 2: File-name based search — ensures documents whose names match query terms
        # are always included (catches large files with low BM25 scores)
        like_terms = expanded_terms(query)
        if like_terms:
            fn_clauses = ["d.file_name LIKE ?"] * len(like_terms[:8])
            fn_params = [f"%{t}%" for t in like_terms[:8]]
            fn_rows = con.execute(
                f"""
                SELECT c.id, d.file_name, c.title, c.content, c.page, c.sheet
                FROM chunks c JOIN documents d ON c.document_id=d.id
                WHERE ({' OR '.join(fn_clauses)})
                ORDER BY d.file_name, c.chunk_index
                LIMIT ?
                """,
                (*fn_params, limit * 4),
            ).fetchall()
            for r in fn_rows:
                if len(hits) >= limit:
                    break
                _add_hit(r, 0.0)

        # Step 3: Content LIKE fallback for remaining slots
        if len(hits) < limit and like_terms:
            ct_clauses = ["c.content LIKE ?"] * len(like_terms[:8])
            ct_params = [f"%{t}%" for t in like_terms[:8]]
            ct_rows = con.execute(
                f"""
                SELECT c.id, d.file_name, c.title, c.content, c.page, c.sheet
                FROM chunks c JOIN documents d ON c.document_id=d.id
                WHERE ({' OR '.join(ct_clauses)})
                ORDER BY d.file_name, c.chunk_index
                LIMIT ?
                """,
                (*ct_params, limit * 4),
            ).fetchall()
            for r in ct_rows:
                if len(hits) >= limit:
                    break
                _add_hit(r, 0.0)
    return hits


def source_list(hits: list[SearchHit]) -> list[str]:
    out = []
    for h in hits:
        if h.file_name not in out:
            out.append(h.file_name)
    return out


def drive_source_map() -> dict[str, str]:
    """Map local/extracted file names to original Google Drive URLs."""
    if not DRIVE_MANIFEST_PATH.exists():
        return {}
    try:
        data = json.loads(DRIVE_MANIFEST_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}
    out: dict[str, str] = {}
    for item in data.get("files", []):
        url = item.get("url")
        if not url:
            continue
        for key in (item.get("local_name"), item.get("extracted_name"), item.get("name")):
            if key:
                out[key] = url
    return out


def format_source(name: str) -> str:
    url = drive_source_map().get(name)
    if not url and name.endswith(".txt"):
        url = drive_source_map().get(name[:-4])
    return f"{name}：{url}" if url else name


def format_sources(names: list[str]) -> list[str]:
    return [format_source(n) for n in names]


def is_sensitive_query(query: str) -> bool:
    return any(term in query for term in SENSITIVE_TERMS)


def truncate_line_text(text: str, max_chars: int = 4800) -> str:
    return text if len(text) <= max_chars else text[: max_chars - 30].rstrip() + "\n……（因 LINE 字數限制節錄）"


def normalize_line_answer(text: str) -> str:
    """Make LLM/Codex output look native in LINE (LINE does not render Markdown)."""
    text = clean_text(text)
    # Remove Markdown markers that LINE shows literally.
    text = text.replace("**", "").replace("__", "")
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^[-*]\s+(?=(卡片摘要|正式報告|資料來源|小秘書建議)[:：])", "", text, flags=re.MULTILINE)
    # The Flex card already provides the card summary; remove duplicate text heading if Codex emits it.
    text = re.sub(r"^卡片摘要[:：]\s*\n?", "", text).strip()
    return truncate_line_text(text)


def list_documents() -> list[str]:
    """Return all indexed document file names."""
    init_db()
    with sqlite3.connect(DB_PATH) as con:
        rows = con.execute("SELECT file_name FROM documents ORDER BY file_name").fetchall()
    return [r[0] for r in rows]


def extract_image_text(image_bytes: bytes) -> str:
    """Use Gemini Vision to extract text from an image (e.g., scanned document photo)."""
    api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
    if not api_key:
        return ""
    try:
        import google.generativeai as genai
        import PIL.Image
        import io
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(os.getenv("GEMINI_MODEL", "gemini-2.5-flash"))
        img = PIL.Image.open(io.BytesIO(image_bytes))
        resp = model.generate_content([
            "請將這張圖片中的所有文字完整辨識並輸出，保持原始格式（段落、表格、清單）。若為公文或表單，請保留欄位名稱與數值的對應關係。",
            img
        ])
        return clean_text(getattr(resp, "text", "") or "")
    except Exception:
        return ""


def _hits_are_relevant(query: str, hits: list[SearchHit]) -> bool:
    """Check if top hits actually contain query terms (prevent hallucination on empty results)."""
    if not hits:
        return False
    query_terms = [t for t in re.findall(r"[一-鿿]{2,}|[A-Za-z0-9]{2,}", query) if len(t) >= 2]
    if not query_terms:
        return True
    top_content = " ".join(h.content + " " + h.file_name for h in hits[:3])
    matched = sum(1 for t in query_terms if t in top_content)
    # Require at least 1 meaningful term to match
    return matched >= 1 and hits[0].score < 0  # bm25 negative = matched; score=0.0 means fallback LIKE hit


def llm_answer(query: str, hits: list[SearchHit], history: list[tuple[str, str]] | None = None) -> str | None:
    api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
    if not api_key or not hits:
        return None
    # Only call LLM if hits are from FTS (score < 0 = BM25 matched), not just fallback LIKE hits
    fts_hits = [h for h in hits if h.score < 0]
    if not fts_hits:
        return None  # No FTS match → skip LLM, use extractive_answer (no-data message)
    try:
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(os.getenv("GEMINI_MODEL", "gemini-2.5-flash"))
        context = "\n\n".join(
            f"【來源：{format_source(h.file_name)}】\n{h.content[:1800]}" for h in fts_hits[:6]
        )
        history_block = ""
        if history:
            recent = history[-3:]
            lines = []
            for q, a in recent:
                lines.append(f"用戶：{q}")
                lines.append(f"小秘書：{a[:300]}…")
            history_block = "\n對話紀錄（最近幾輪，供參考）：\n" + "\n".join(lines) + "\n"
        prompt = f"""你是「二大隊行政小秘書」，協助臺北市消防局第二救災救護大隊依據資料庫內容回覆議會諮詢。

【絕對禁止事項】
- 禁止在「資料庫片段」中找不到相關數字、日期、場所時自行捏造任何數據
- 禁止憑常識或推測補充資料庫片段沒有的內容
- 若片段內容與問題不相關，必須直接回覆「現有資料庫未見明確內容，請確認相關文件已上傳」
- 不得生成假統計表、假清冊、假數字

硬性規則：
1. 只能根據下方「資料庫片段」原文回答。
2. 資料不足時輸出：「現有資料庫未見明確內容，請確認相關文件已上傳並重建索引。」
3. 用正式公務語氣。
4. 敏感資料（人員名冊、危險物、毒化物）只提供統計摘要，不揭露完整明細。
5. 回覆最後列「資料來源」，只列實際找到資料的檔案；若有雲端連結，必須列原始 Google Drive 連結，不要只列本機檔名。
{history_block}
使用者問題：{query}

資料庫片段（這是唯一可用的資料，不得超出此範圍回答）：
{context}

請依此格式回答：
報告人：二大隊行政小秘書
日期：{taiwan_minguo_date()}

根據……說明如下：

一、查詢結果摘要：

二、相關資料明細：

三、深度分析與小秘書建議：

資料來源：
- ...

以上報告。"""
        resp = model.generate_content(prompt)
        text = getattr(resp, "text", "") or ""
        return clean_text(text)
    except Exception:
        return None


def extractive_answer(query: str, hits: list[SearchHit]) -> str:
    if not hits:
        return f"""報告人：二大隊行政小秘書
日期：{taiwan_minguo_date()}

經查目前資料庫內未見與「{query}」直接相關之明確內容。

小秘書建議：
一、請確認議員備詢資料是否已完成上傳及重建索引。
二、如屬最新業務數據，建議洽承辦組別確認後再行答覆。

以上報告。"""
    sources = source_list(hits)
    snippets = []
    for i, h in enumerate(hits[:5], start=1):
        snippet = h.content.strip().replace("\n", "\n   ")
        snippets.append(f"{i}. 【{h.file_name}】\n   {snippet[:700]}")
    source_lines = format_sources(sources)
    return f"""報告人：二大隊行政小秘書
日期：{taiwan_minguo_date()}

根據資料庫檢索結果，針對「{query}」說明如下：

一、查詢結果摘要：
已於議員備詢資料庫中查得相關內容，主要來源包含：{ '、'.join(source_lines[:5]) }。

二、相關資料明細：
{chr(10).join(snippets)}

三、深度分析與小秘書建議：
1. 建議以上開來源文件作為答覆依據，必要時再由承辦組別確認最新數字。
2. 若本題涉及個資、救災部署、危險物場所或其他敏感資料，建議對外回覆時以統計摘要及管理作為為主，避免直接揭露完整明細。
3. 如需議員書面答覆版本，可直接追問「改寫成議員書面答覆」。

資料來源：
{chr(10).join('- ' + s for s in source_lines)}

以上報告。"""


def _hits_match_query(query: str, hits: list[SearchHit]) -> bool:
    """
    Return True only if the hits contain actual subject terms from the query.
    Numbers/years alone do NOT count as evidence of relevance.
    """
    if not hits:
        return False
    # Only check content from NON-Drive-list files (Drive list has file names that cause false positives)
    real_content = " ".join(h.content for h in hits if _DRIVE_LIST_KEYWORD not in h.file_name)
    if not real_content:
        return False  # All hits are from Drive list → handled by _only_drive_list_hits
    def _count(term: str, text: str) -> int:
        return text.count(term)

    # 1. Check KNOWN_TERMS in real content (require ≥2 occurrences for weak matches)
    for t in KNOWN_TERMS:
        if t in query and t in real_content:
            if _count(t, real_content) >= 2:
                return True
    # 2. Check SYNONYMS in real content
    for k, vals in SYNONYMS.items():
        if k in query:
            for term in [k] + vals:
                if _count(term, real_content) >= 2:
                    return True
    # 3. Extract bigrams/trigrams from query and check in real content (require ≥2 occurrences)
    question_stop = {"什麼", "哪些", "有什", "沒有", "如何", "怎樣", "幾個", "多少",
                     "請問", "告訴", "想知", "知道", "查詢", "有關", "情形", "情況", "說明"}
    cn_only = re.sub(r"[^一-鿿]", "", query)
    candidate_terms: set[str] = set()
    for n in (4, 3, 2):
        for i in range(max(0, len(cn_only) - n + 1)):
            gram = cn_only[i:i+n]
            if gram not in question_stop:
                candidate_terms.add(gram)
    content_matches = [t for t in candidate_terms if len(t) >= 2 and _count(t, real_content) >= 2]
    if content_matches:
        return True
    return False


_DRIVE_LIST_KEYWORD = "Drive檔案清單"


def _only_drive_list_hits(hits: list[SearchHit]) -> bool:
    """True if all hits come solely from the Drive file index (no actual content files)."""
    return bool(hits) and all(_DRIVE_LIST_KEYWORD in h.file_name for h in hits)


def _drive_list_has_related_file(query: str, hits: list[SearchHit]) -> list[str]:
    """Return file names mentioned in the Drive list that seem related to the query."""
    terms = [t for t in KNOWN_TERMS if t in query] or [
        cn_only[i:i+2] for cn_only in [re.sub(r"[^一-鿿]", "", query)]
        for i in range(len(cn_only) - 1)
    ]
    related = []
    for h in hits:
        if _DRIVE_LIST_KEYWORD not in h.file_name:
            continue
        for line in h.content.split("\n"):
            if any(t in line for t in terms) and line.strip().startswith(("1", "2", "3", "4", "5", "6", "7", "8", "9", "-", "·")):
                name = re.sub(r"^\d+\.\s*", "", line.strip())[:60]
                if name and name not in related:
                    related.append(name)
    return related[:5]


def codex_answer(query: str, history: list[tuple[str, str]] | None = None) -> str | None:
    """Use Codex CLI as the answering agent for LINE test mode.

    This intentionally lets Codex inspect the project/data folder in read-only mode.
    Returns None on failure so the caller can fall back to the normal RAG path.
    """
    codex_bin = os.getenv("CODEX_BIN", "/opt/homebrew/bin/codex")
    model = os.getenv("CODEX_MODEL", "gpt-5.4")
    timeout = int(os.getenv("CODEX_TIMEOUT", "150"))
    history_block = ""
    if history:
        recent = history[-3:]
        history_block = "\n最近對話紀錄：\n" + "\n".join(
            f"使用者：{q}\n小秘書：{a[:500]}" for q, a in recent
        )

    prompt = f"""你現在是「二大隊行政小秘書」的 LINE 回覆代理。請直接根據目前工作資料夾回覆使用者問題。

工作根目錄：{BASE_DIR}
主要資料位置：
- 原始檔：{RAW_DIR}
- 抽取文字：{EXTRACTED_DIR}
- SQLite 索引：{DB_PATH}
- Google Drive 原始連結對照表：{DRIVE_MANIFEST_PATH}

請遵守：
1. 只能讀取資料，不得修改、新增或刪除任何檔案。
2. 優先查 data/extracted、data/raw、data/index.sqlite3 內與問題相關的資料；必要時可用 Python/SQLite 查詢。
3. LINE 會另外顯示 Flex 卡片摘要；你的文字訊息不要再輸出「卡片摘要」段落，請直接從正式報告正文開始。
4. 文字格式必須像正式公務報告：開頭「根據……說明如下：」，中段用「一、二、三、四、五、」分節，最後列「資料來源」與「小秘書建議」，結尾固定「以上報告。」
5. 絕對不要使用 Markdown 符號，例如 **粗體**、## 標題、```；LINE 會把星號原樣顯示，造成版面很醜。
6. 若資料不足，明確寫「現有資料庫未見明確內容」，不得捏造數字、日期、地點或人名。
7. 涉及人員名冊、個資、危險物、毒化物、救災部署等敏感資料，只提供摘要與建議，不揭露完整明細。
8. 資料來源請列實際查到的原始雲端檔案連結；請用 Google Drive 原始連結對照表把本機檔名轉成雲端 URL，不要只列本機 .txt 檔名。
9. 回覆要比一般摘要更像「深度報告」：若資料足夠，請包含現況、分布/統計、重點分析、風險或管理意涵、建議作法；但仍需控制在 LINE 可閱讀長度。
10. 不要描述你執行了哪些 shell 指令。
{history_block}
使用者問題：{query}
"""

    with tempfile.NamedTemporaryFile("w+", suffix=".txt", delete=False) as f:
        out_path = f.name
    import time as _time
    t0 = _time.monotonic()
    try:
        cmd = [
            codex_bin,
            "exec",
            "--skip-git-repo-check",
            "--sandbox",
            "read-only",
            "-C",
            str(BASE_DIR),
            "-m",
            model,
            "-o",
            out_path,
            prompt,
        ]
        env = os.environ.copy()
        env.setdefault("HOME", str(Path.home()))
        proc = subprocess.run(
            cmd,
            input="",
            text=True,
            capture_output=True,
            timeout=timeout,
            env=env,
        )
        elapsed = _time.monotonic() - t0
        print(f"[Codex] query={query[:60]!r} model={model} elapsed={elapsed:.1f}s returncode={proc.returncode}", flush=True)
        if proc.returncode != 0:
            return None
        text = Path(out_path).read_text(encoding="utf-8", errors="ignore").strip()
        return normalize_line_answer(text) if text else None
    except Exception as e:
        elapsed = _time.monotonic() - t0
        print(f"[Codex] failed after {elapsed:.1f}s: {e}", flush=True)
        return None
    finally:
        try:
            Path(out_path).unlink(missing_ok=True)
        except Exception:
            pass


def answer(query: str, history: list[tuple[str, str]] | None = None) -> str:
    hits = search(query, limit=8)

    # Case 1: Drive file list is the ONLY source of hits
    if _only_drive_list_hits(hits):
        related = _drive_list_has_related_file(query, hits)
        if related:
            file_list = "\n".join(f"  - {f}" for f in related)
            return truncate_line_text(f"""報告人：二大隊行政小秘書
日期：{taiwan_minguo_date()}

經查 Drive 檔案清單，與「{query}」相關之原始文件索引如下：
{file_list}

⚠️ 小秘書說明：
以上為 Drive 雲端資料夾中已知存在的相關檔案名稱，但原始文件尚未下載至本機資料庫，因此無法提供具體數字或內容。

建議：
一、請將上述原始 Word/Excel 檔案下載後放入 data/raw/ 資料夾。
二、在 LINE 傳送 /重建索引 完成更新。
三、重建後即可得到詳細的正式報告。

以上報告。""")
        return truncate_line_text(f"""報告人：二大隊行政小秘書
日期：{taiwan_minguo_date()}

經查目前資料庫內未見與「{query}」直接相關之明確內容。

小秘書建議：
一、請確認相關文件是否已上傳至 data/raw/ 資料夾。
二、上傳後請執行 /重建索引 更新資料庫。
三、如屬最新業務數據，建議洽承辦組別確認後再行答覆。

以上報告。""")

    # Case 2: No relevant content match
    if not _hits_match_query(query, hits):
        return truncate_line_text(f"""報告人：二大隊行政小秘書
日期：{taiwan_minguo_date()}

經查目前資料庫內未見與「{query}」直接相關之明確內容。

小秘書建議：
一、請確認相關文件是否已上傳至 data/raw/ 資料夾。
二、上傳後請執行 /重建索引 更新資料庫。
三、如屬最新業務數據，建議洽承辦組別確認後再行答覆。

以上報告。""")

    # Case 3: Actual content available
    text = llm_answer(query, hits, history) or extractive_answer(query, hits)
    return normalize_line_answer(text)
